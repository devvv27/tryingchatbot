from __future__ import annotations

from pathlib import Path
from typing import Any

import requests
from bs4 import BeautifulSoup
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from app.config import CHUNK_OVERLAP, CHUNK_SIZE

try:
    from pdf2image import convert_from_path  # type: ignore
    import pytesseract  # type: ignore

    _OCR_AVAILABLE = True
except Exception:
    convert_from_path = None  # type: ignore[assignment]
    pytesseract = None  # type: ignore[assignment]
    _OCR_AVAILABLE = False


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    text = " ".join(text.split())
    if not text:
        return []

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunks.append(text[start:end])
        if end == len(text):
            break
        start = max(end - overlap, start + 1)
    return chunks


def _ocr_pdf_pages(file_path: Path) -> list[str]:
    if not _OCR_AVAILABLE or convert_from_path is None or pytesseract is None:
        return []

    try:
        images = convert_from_path(str(file_path), dpi=200)
    except Exception:
        return []

    page_texts: list[str] = []
    for image in images:
        try:
            page_texts.append(pytesseract.image_to_string(image) or "")
        except Exception:
            page_texts.append("")
    return page_texts


def parse_pdf(file_path: Path) -> list[dict[str, Any]]:
    reader = PdfReader(str(file_path))
    records: list[dict[str, Any]] = []
    ocr_texts: list[str] | None = None

    needs_ocr = any(not (page.extract_text() or "").strip() for page in reader.pages)
    if needs_ocr:
        ocr_texts = _ocr_pdf_pages(file_path)

    for page_idx, page in enumerate(reader.pages, start=1):
        raw = page.extract_text() or ""
        if not raw.strip() and ocr_texts and len(ocr_texts) >= page_idx:
            raw = ocr_texts[page_idx - 1]
        for chunk_idx, chunk in enumerate(chunk_text(raw)):
            records.append(
                {
                    "text": chunk,
                    "quote_text": chunk,
                    "location_type": "page",
                    "location_value": f"Page {page_idx}",
                    "metadata": {"page": page_idx, "chunk_index": chunk_idx},
                }
            )
    return records


def parse_docx(file_path: Path) -> list[dict[str, Any]]:
    doc = Document(str(file_path))
    sections: list[tuple[str, str]] = []
    current_heading = "Body"
    buffer: list[str] = []

    def flush() -> None:
        if buffer:
            sections.append((current_heading, "\n".join(buffer)))
            buffer.clear()

    for p in doc.paragraphs:
        value = p.text.strip()
        if not value:
            continue
        style_name = (p.style.name or "").lower()
        if "heading" in style_name:
            flush()
            current_heading = value
        else:
            buffer.append(value)

    # Many assignment documents store questions in tables; index table text as section content.
    for table_idx, table in enumerate(doc.tables, start=1):
        for row_idx, row in enumerate(table.rows, start=1):
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if not cells:
                continue
            buffer.append(f"Table {table_idx} Row {row_idx}: {' | '.join(cells)}")
    flush()

    records: list[dict[str, Any]] = []
    for section_name, section_text in sections:
        for chunk_idx, chunk in enumerate(chunk_text(section_text)):
            records.append(
                {
                    "text": chunk,
                    "quote_text": chunk,
                    "location_type": "section",
                    "location_value": f"Section: {section_name}",
                    "metadata": {"section": section_name, "chunk_index": chunk_idx},
                }
            )
    return records


def parse_xlsx(file_path: Path) -> list[dict[str, Any]]:
    wb = load_workbook(str(file_path), data_only=True, read_only=True)
    records: list[dict[str, Any]] = []

    for sheet in wb.worksheets:
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            cells = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if not cells:
                continue
            row_text = " | ".join(cells)
            records.append(
                {
                    "text": row_text,
                    "quote_text": row_text,
                    "location_type": "tab_row",
                    "location_value": f"Tab: {sheet.title}, Row: {row_idx}",
                    "metadata": {"sheet": sheet.title, "row": row_idx},
                }
            )
    return records


def parse_url(url: str) -> list[dict[str, Any]]:
    response = requests.get(url, timeout=20)
    response.raise_for_status()

    soup = BeautifulSoup(response.text, "html.parser")
    title = soup.title.get_text(" ", strip=True) if soup.title else "Web Page"

    # Collect content under heading tags to preserve section-like navigation.
    sections: list[tuple[str, list[str]]] = []
    current_heading = title
    current_paragraphs: list[str] = []

    for el in soup.find_all(["h1", "h2", "h3", "p", "li"]):
        tag = el.name.lower()
        text = el.get_text(" ", strip=True)
        if not text:
            continue
        if tag.startswith("h"):
            if current_paragraphs:
                sections.append((current_heading, current_paragraphs))
            current_heading = text
            current_paragraphs = []
        else:
            current_paragraphs.append(text)

    if current_paragraphs:
        sections.append((current_heading, current_paragraphs))

    records: list[dict[str, Any]] = []
    for heading, paragraphs in sections:
        merged = "\n".join(paragraphs)
        for chunk_idx, chunk in enumerate(chunk_text(merged)):
            records.append(
                {
                    "text": chunk,
                    "quote_text": chunk,
                    "location_type": "section",
                    "location_value": f"Section: {heading}",
                    "metadata": {"url": url, "section": heading, "chunk_index": chunk_idx},
                }
            )
    return records


def parse_source(source_type: str, source_path: Path | None = None, source_url: str | None = None) -> list[dict[str, Any]]:
    if source_type == "pdf" and source_path:
        return parse_pdf(source_path)
    if source_type == "docx" and source_path:
        return parse_docx(source_path)
    if source_type == "xlsx" and source_path:
        return parse_xlsx(source_path)
    if source_type == "url" and source_url:
        return parse_url(source_url)
    raise ValueError("Unsupported source input")
